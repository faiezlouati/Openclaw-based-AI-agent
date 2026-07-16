#!/usr/bin/env python3

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY        = RGBColor(0x1F, 0x38, 0x64)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
HIGH_RED    = RGBColor(0xC0, 0x00, 0x00)
MED_ORANGE  = RGBColor(0xFF, 0x66, 0x00)
LOW_GREEN   = RGBColor(0x37, 0x56, 0x23)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
BODY_GREY   = RGBColor(0x40, 0x40, 0x40)



# Handles shade cell.
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.upper().replace('#', ''))
    tcPr.append(shd)

# Handles add cell borders.
def add_cell_borders(cell, color='CCCCCC', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)

# Handles add table row.
def add_table_row(table, values, header=False, severity_col=None, alert_col=None):
    row = table.add_row()
    for ci, val in enumerate(values):
        cell = row.cells[ci]
        cell.text = ''

        is_severity = (severity_col is not None and ci == severity_col)
        is_alert    = (alert_col is not None and ci == alert_col)

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)

        run = para.add_run(val)
        run.font.size = Pt(9)

        if header:
            run.bold = True
            run.font.color.rgb = BLACK
            shade_cell(cell, 'FFFFFF')
            add_cell_borders(cell, color='2E75B6', sz=6)
        elif is_severity or is_alert:
            run.bold = True
            v = val.strip().upper()
            if 'HIGH' in v or (val and val[0] == '🔴'):
                run.font.color.rgb = HIGH_RED
            elif 'MEDIUM' in v or (val and val[0] == '🟡'):
                run.font.color.rgb = MED_ORANGE
            elif 'LOW' in v or (val and val[0] == '🟢'):
                run.font.color.rgb = LOW_GREEN
            shade_cell(cell, 'FFFFFF')
            add_cell_borders(cell, color='CCCCCC', sz=4)
        else:
            run.font.color.rgb = BLACK
            shade_cell(cell, 'FFFFFF')
            add_cell_borders(cell, color='CCCCCC', sz=4)
    return row

# Handles section heading.
def section_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(13)
    return para

# Handles body para.
def body_para(doc, text, size=9):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BODY_GREY
    return para

# Handles add callout border.
def add_callout_border(para, color='1F3864', sz=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    str(sz))
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), color)
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EEF3FA')
    pPr.append(shd)



# Handles make header section.
def make_header_section(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(10)
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(11)
        r2.font.color.rgb = ACCENT_BLUE



# Builds ccap docx.
def build_ccap_docx(tender_ref, src_md, out_docx):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    make_header_section(doc,
        f"CCAP RELEVANCE ANALYSIS — {tender_ref}",
        "Administrative Tender Document | Ministry of National Defence")


    section_heading(doc, "1. Document Overview")

    t1 = doc.add_table(rows=1, cols=2)
    t1.style = 'Table Grid'
    add_table_row(t1, ['Field', 'Detail'], header=True)
    rows1 = [
        ('Authority', 'Ministry of National Defence — General Directorate for ICT'),
        ('Project', 'Acquisition and installation of integrated servers'),
        ('Tender Reference', 'N° 06/DIV/2026'),
        ('Document Type', 'Cahier des Charges Administratives (CCAP)'),
        ('Language', 'Arabic (original); English translation'),
        ('Currency', 'Tunisian Dinar (TND)'),
        ('Bid Validity', '120 days from submission deadline'),
        ('Temporary Guarantee', '5,000 TND'),
    ]
    for r in rows1:
        add_table_row(t1, list(r))
    doc.add_paragraph('')


    section_heading(doc, "2. Participation Conditions")

    body_para(doc, "Article 2 — Scope", bold=True) if False else None
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("Article 2 — Scope")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY

    for txt in ["Single lot consisting of 8 line items",
                "Bidders must submit for the entire lot — partial bids are NOT accepted"]:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(1)
        run = para.add_run(txt)
        run.font.size = Pt(9); run.font.color.rgb = BODY_GREY

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("Article 3 — Eligibility")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY

    for txt in ["Participating suppliers must demonstrate professional, technical, and financial qualifications"]:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(1)
        run = para.add_run(txt)
        run.font.size = Pt(9); run.font.color.rgb = BODY_GREY

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("Article 4 — Access to Documents")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY

    for txt in ["Complete tender package available via TUNEPS online portal upon publication"]:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(1)
        run = para.add_run(txt)
        run.font.size = Pt(9); run.font.color.rgb = BODY_GREY
    doc.add_paragraph('')


    section_heading(doc, "3. Bid Components")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("3.1 Financial Offer")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY

    t3a = doc.add_table(rows=1, cols=2)
    t3a.style = 'Table Grid'
    add_table_row(t3a, ['Document', 'Requirements'], header=True)
    add_table_row(t3a, ['Financial Commitment Letter', 'Completed, signed, stamped per Appendix 1 template'])
    add_table_row(t3a, ['Detailed Estimative Invoice', 'Itemized per component, signed and stamped'])
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("3.2 Technical Offer")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY

    t3b = doc.add_table(rows=1, cols=2)
    t3b.style = 'Table Grid'
    add_table_row(t3b, ['Document', 'Requirements'], header=True)
    add_table_row(t3b, ['Signed CCTP', 'Initialed every page; signed and stamped on final page'])
    add_table_row(t3b, ['Technical Bid Forms', 'Completed per templates in technical specifications'])
    add_table_row(t3b, ['Technical Data Sheets', 'Completed per lot item, signed and stamped'])
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("3.3 Administrative Documents")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY

    t3c = doc.add_table(rows=1, cols=2)
    t3c.style = 'Table Grid'
    add_table_row(t3c, ['Document', 'Requirements'], header=True)
    add_table_row(t3c, ['Temporary Financial Guarantee', 'Bank guarantee or solidarity undertaking; 120-day validity'])
    add_table_row(t3c, ['Declaration of Non-Collusion', 'Sworn statement via TUNEPS portal'])
    add_table_row(t3c, ['Declaration — No Conflict of Interest', 'Sworn statement via TUNEPS portal'])
    add_table_row(t3c, ['Manager Identification Card', 'Civil status card + national ID copy + 2 photos'])
    add_table_row(t3c, ['Company Registry Extract', 'Within 3 months of issue date'])
    doc.add_paragraph('')


    section_heading(doc, "4. Financial Guarantee (Article 10)")

    t4 = doc.add_table(rows=1, cols=2)
    t4.style = 'Table Grid'
    add_table_row(t4, ['Element', 'Detail'], header=True)
    add_table_row(t4, ['Amount', '5,000 TND'])
    add_table_row(t4, ['Form', 'Bank guarantee or solidarity liability undertaking'])
    add_table_row(t4, ['Validity', 'Must cover full bid validity period (minimum 120 days)'])
    add_table_row(t4, ['Release — Non-selected bidders', 'Within 20 days of contract award notification'])
    add_table_row(t4, ['Release — Selected bidder', 'Upon submission of final performance guarantee'])
    add_table_row(t4, ['Release — Automatically rejected', 'Immediate release'])
    add_table_row(t4, ['Seizure — Withdrawal during validity', 'Bid withdrawal after submission deadline'])
    add_table_row(t4, ['Seizure — Contract refusal', 'Selected bidder refuses to sign within deadline'])
    add_table_row(t4, ['Seizure — No final guarantee', 'Selected bidder fails to provide final performance guarantee'])
    doc.add_paragraph('')


    section_heading(doc, "5. Submission Rules (Article 8)")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("Online Submission (via TUNEPS)")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY

    for txt in [
        "Registration in TUNEPS is MANDATORY before the submission deadline",
        "Financial and technical offers submitted via TUNEPS portal"
    ]:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(1)
        run = para.add_run(txt)
        run.font.size = Pt(9); run.font.color.rgb = BODY_GREY

    for label, items in [
        ("Late Submissions", ["Bids received or modified after the deadline are AUTOMATICALLY REJECTED — no exceptions"]),
        ("Modifications", [
            "Modifications permitted during the submission window",
            "After deadline: no modification, correction, addition, or withdrawal is allowed",
            "Bids containing additions/modifications must identify altered items by signing person"
        ]),
        ("Oversized Technical Bids", [
            "If online submission volume is exceeded, remaining technical documents may be sent offline",
            "Offline submission list must be declared in the electronic bid"
        ]),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(label)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY
        for txt in items:
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            run = para.add_run(txt)
            run.font.size = Pt(9); run.font.color.rgb = BODY_GREY
    doc.add_paragraph('')


    section_heading(doc, "6. Evaluation (Article 13)")

    t6 = doc.add_table(rows=1, cols=2)
    t6.style = 'Table Grid'
    add_table_row(t6, ['Phase', 'Activity'], header=True)
    add_table_row(t6, ['Stage 1', 'Administrative compliance + financial guarantee verification'])
    add_table_row(t6, ['Stage 2', 'Financial offer correction — arithmetic/material errors corrected; ascending price ranking'])
    add_table_row(t6, ['Stage 3', 'Lowest evaluated compliant bid wins (Moins-disant)'])
    doc.add_paragraph('')


    section_heading(doc, "7. Key Risks & Observations")

    t7 = doc.add_table(rows=1, cols=3)
    t7.style = 'Table Grid'
    add_table_row(t7, ['#', 'Risk / Observation', 'Severity'], header=True)
    risks = [
        ('1', '120-day bid validity is relatively long — bidders must maintain pricing for 4 months', 'Medium'),
        ('2', 'Single-lot / 8 items — no partial bid option increases risk for specialist suppliers', 'High'),
        ('3', '5,000 TND temporary guarantee (0.1% of 5M project) is unusually low — potential budget constraint or risk allocation preference', 'Medium'),
        ('4', 'Out-of-line document submission explicitly permitted — creates administrative complexity', 'Medium'),
        ('5', 'Non-public bid opening session — reduced transparency in evaluation', 'Medium'),
        ('6', 'BCT authorization required for foreign currency bids — procedural burden', 'Low'),
        ('7', '20-day guarantee release for non-selected bidders — reasonable', 'Low'),
        ('8', '7-day cure period for missing documents after opening — unusual, potentially bidder-friendly', 'Low'),
    ]
    for r in risks:
        add_table_row(t7, list(r), severity_col=2)
    doc.add_paragraph('')


    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after  = Pt(6)
    add_callout_border(callout, color='1F3864', sz=6)
    run = callout.add_run(
        "This CCAP analysis is for informational purposes. Bidders should verify all conditions "
        "against the official Arabic CCAP document before submitting.")
    run.font.size = Pt(9)
    run.font.color.rgb = BODY_GREY


    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_f.add_run(f"CCAP Analysis | Tender Ref: {tender_ref} | Generated by Tuneps Analyst")
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run_f.italic = True

    doc.save(out_docx)
    print(f"CCAP DOCX saved: {out_docx}")

if __name__ == '__main__':
    tender_ref = sys.argv[1] if len(sys.argv) > 1 else '20260301601'
    out_dir   = Path(sys.argv[2] if len(sys.argv) > 2 else '.').resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    out_docx  = out_dir / f"CCAP_Analysis_{tender_ref}.docx"
    build_ccap_docx(tender_ref, None, str(out_docx))

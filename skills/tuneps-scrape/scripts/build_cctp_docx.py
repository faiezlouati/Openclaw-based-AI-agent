#!/usr/bin/env python3

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


NAVY        = RGBColor(0x1F, 0x38, 0x64)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
HIGH_RED    = RGBColor(0xC0, 0x00, 0x00)
MED_ORANGE  = RGBColor(0xFF, 0x66, 0x00)
LOW_GREEN   = RGBColor(0x37, 0x56, 0x23)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
BODY_GREY   = RGBColor(0x40, 0x40, 0x40)
LIGHT_BG    = 'EEF3FA'



# Handles shade cell.
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.upper().replace('#', ''))
    tcPr.append(shd)

# Handles set cell border.
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        if side in kwargs:
            color, sz = kwargs[side]
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    str(sz))
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), color)
            tcBorders.append(b)
    tcPr.append(tcBorders)

# Handles add cell borders.
def add_cell_borders(cell, color='AAAAAA', sz=4):
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        tcBorders.append(b)

# Handles add table row.
def add_table_row(table, values, header=False, risk_col=None):
    row = table.add_row()
    for ci, val in enumerate(values):
        cell = row.cells[ci]
        cell.text = ''
        is_severity = (risk_col is not None and ci == risk_col)

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)

        run = para.add_run(val)
        run.font.size = Pt(9)

        if header:
            run.bold = True
            run.font.color.rgb = BLACK
            shade_cell(cell, 'FFFFFF')
            add_cell_borders(cell, color='2E75B6', sz=6)
        elif is_severity:
            run.bold = True
            v = val.strip().upper()
            if 'HIGH' in v:
                run.font.color.rgb = HIGH_RED
            elif 'MEDIUM' in v:
                run.font.color.rgb = MED_ORANGE
            elif 'LOW' in v:
                run.font.color.rgb = LOW_GREEN
            shade_cell(cell, 'FFFFFF')
            add_cell_borders(cell, color='CCCCCC', sz=4)
        else:
            run.font.color.rgb = BLACK
            shade_cell(cell, 'FFFFFF')
            add_cell_borders(cell, color='CCCCCC', sz=4)
    return row

# Handles add paragraph border.
def add_paragraph_border(para, color='1F3864', sz=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    str(sz))
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), color)
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  LIGHT_BG)
    pPr.append(shd)



doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)



title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(0)
title_para.paragraph_format.space_after  = Pt(4)
run = title_para.add_run('CCTP RELEVANCE ANALYSIS')
run.bold = True
run.font.color.rgb = NAVY
run.font.size = Pt(16)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_before = Pt(0)
sub_para.paragraph_format.space_after  = Pt(10)
run = sub_para.add_run('Datacenter Modernization — Mutuelle de l\'Armée Nationale (Tunisia)')
run.font.color.rgb = ACCENT_BLUE
run.font.size = Pt(11)



# Handles section heading.
def section_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(4)

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(13)
    return para

# Handles body para.
def body_para(doc, text, size=10):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BODY_GREY
    return para

# Handles bullet para.
def bullet_para(doc, text, size=10):
    para = doc.add_paragraph(style='List Paragraph')
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BODY_GREY
    return para




section_heading(doc, '1. Extraction and Mapping')

body_para(doc,
    'The Mutuelle de l\'Armée Nationale, operating under the supervision of the Ministry of National '
    'Defense of Tunisia, has issued a public tender for the complete modernization of its central '
    'datacenter infrastructure. The stated goal of the project is to strengthen service continuity '
    'for hosted applications, optimize resource management, and renew the virtualization platform '
    'that currently runs on VMware vSphere 8.')

body_para(doc,
    'The tender takes the form of a single contract composed of eight indissociable items covering '
    'compute, storage, backup, networking, security, and physical infrastructure. The successful '
    'bidder must deliver a turnkey solution that includes hardware, software, services, and post-sale '
    'support. The execution period is 120 days from contract approval. The existing VMware vSphere 8 '
    'platform must be migrated to the new solution as part of the project scope.')


items_table = doc.add_table(rows=0, cols=4)
items_table.style = 'Table Grid'
add_table_row(items_table, ['#', 'Item', 'Qty', 'Category'], header=True)
items_data = [
    ('1', 'Virtualization servers',       '3', 'Compute'),
    ('2', 'Virtualization software',      '1', 'Hypervisor'),
    ('3', 'Storage array',                '1', 'Storage'),
    ('4', 'Backup server',                '1', 'Compute'),
    ('5', 'Backup software',              '1', 'Data protection'),
    ('6', 'Datacenter switches',          '2', 'Network'),
    ('7', 'Firewall appliances',          '2', 'Security'),
    ('8', 'Rack cabinet 42U',             '1', 'Infrastructure'),
]
for row in items_data:
    add_table_row(items_table, list(row))

doc.add_paragraph('')


intro2 = doc.add_paragraph()
intro2.paragraph_format.space_before = Pt(2)
intro2.paragraph_format.space_after  = Pt(2)
r = intro2.add_run('Included services:')
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = BODY_GREY

body_para(doc,
    'Beyond the supply of equipment and software, the successful bidder is responsible for the '
    'complete deployment and operational handover of the new datacenter platform. '
    'The scope of services includes the following:')

services = [
    'Physical installation, racking, and cabling of all hardware components in the client\'s premises',
    'Configuration of servers, storage array, hypervisor cluster, switches, firewalls, and backup environment',
    'Live migration of all virtual machines and associated data from the existing VMware vSphere 8 platform '
    'to the new hypervisor, without service interruption',
    'Functional and performance testing of the integrated solution under the client\'s supervision before go-live',
    'Two certifying trainings of five working days each, for six staff members per session, delivered in an '
    'accredited training center by certified trainers — one training on the proposed virtualization solution, '
    'the other on the proposed backup solution',
    'Complete technical documentation in French or English, including installation procedures, configuration '
    'files, operating procedures (backup, restore, startup, shutdown), and a configuration dossier that must '
    'be maintained throughout the project lifecycle',
    'Knowledge transfer to the client\'s technical teams during all interventions',
    'Two years of 24/7 manufacturer support, including firmware updates, software patches, and security corrections',
]
for s in services:
    bullet_para(doc, s)

doc.add_paragraph('')




section_heading(doc, '2. Qualification Requirements')

body_para(doc,
    'Technical and product-related conditions defined in the CCTP that a bidder must satisfy to be '
    'eligible. Failure on any line means automatic disqualification at the technical evaluation stage.')

qual_table = doc.add_table(rows=0, cols=2)
qual_table.style = 'Table Grid'
add_table_row(qual_table, ['Category', 'Requirement'], header=True)

qual_data = [
    ('Manufacturer authorisation',
     'MAF letter required for servers and storage array'),
    ('Brand constraint',
     'Servers, storage, and backup server must be from the same manufacturer'),
    ('Firewall vendor',
     'Must be listed in latest Gartner Magic Quadrant for Network Firewalls'),
    ('Firewall certification',
     'Common Criteria EAL 4+ or ICSA Labs'),
    ('Hardware certifications',
     'ISO 9001:2015 and EN 62368, EN 55032, EN 55035 (or equivalent)'),
    ('Training capability',
     'Accredited training center with certified trainers (CVs to be provided)'),
    ('Support capability',
     '2-year manufacturer support, 24/7, on all hardware and software'),
    ('Documentation',
     'Full technical documentation in French or English, free updates'),
]
for cat, req in qual_data:
    add_table_row(qual_table, [cat, req])

doc.add_paragraph('')




section_heading(doc, '3. Technical Demand Summary')

body_para(doc,
    'Headline technical demands per item. Detailed specifications are in the CCTP source document.')

tech_table = doc.add_table(rows=0, cols=2)
tech_table.style = 'Table Grid'
add_table_row(tech_table, ['Item', 'Key Demands'], header=True)

tech_data = [
    ('Virtualization servers (×3)',
     '2U latest generation, 2×16 cores @ 2.4 GHz, 512 GB DDR5, hardware RAID with non-volatile cache, '
     '4×1GbE + 4×10/25GbE, redundant PSU'),
    ('Virtualization software',
     'Bare-metal hypervisor, 3-node HA cluster, VMware vSphere 8 migration tools, boot on SAN, '
     'multi-OS support'),
    ('Storage array',
     'Same brand as servers, dual active/active controllers, 128 GB cache, unified block + file, '
     'all-NVMe SSD, sync + async replication, auto-tiering'),
    ('Backup server',
     'Same brand as virtualization servers, single-CPU 16 cores, 128 GB DDR5, 4×8 TB HDD'),
    ('Backup software',
     'Subscription license 15 VMs (2 years), agentless, image-based, encrypted, app-aware restore, '
     'replication, tape support'),
    ('Datacenter switches (×2)',
     'L3 Top-of-Rack, 600 Gbps switching, 4000 VLANs, 16×1GbE PoE + 4×10GbE SFP+, full L2/L3 + '
     'security features'),
    ('Firewall (×2)',
     '5 Gbps firewall throughput, 4 Gbps IPSec VPN, NGFW (IPS, AV, sandboxing, web/DNS filtering), '
     'HA Active-Active and Active-Passive'),
    ('Rack cabinet',
     '42U, static load ≥1000 kg, front-to-back airflow, 2 intelligent PDUs (redundancy A/B)'),
]
for item, demand in tech_data:
    add_table_row(tech_table, [item, demand])

doc.add_paragraph('')




section_heading(doc, '4. Main Technical Risks')

body_para(doc,
    'Principal technical execution risks identified from the CCTP content.')

risks_table = doc.add_table(rows=0, cols=2)
risks_table.style = 'Table Grid'
add_table_row(risks_table, ['Risk', 'Severity'], header=True)

risks = [
    ('Live VM migration from VMware vSphere 8 to new hypervisor without production impact', 'High'),
    ('Same-brand constraint on servers, storage, and backup server limits sourcing flexibility', 'Medium'),
    ('Backup licence sized at 15 VMs only — actual production VM count not disclosed', 'Medium'),
    ('RAID 5 set as the active level on enterprise all-NVMe storage', 'Medium'),
    ('Battery autonomy on storage array not quantified', 'Medium'),
    ('RPO/RTO targets not defined by the buyer', 'Medium'),
    ('Boot-on-SAN dependency adds configuration complexity', 'Low'),
    ('Sync + async replication on full capacity often licensed separately by OEMs', 'Medium'),
]
for risk, severity in risks:
    row = risks_table.add_row()
    for ci, val in enumerate([risk, severity]):
        cell = row.cells[ci]
        cell.text = ''
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run(val)
        run.font.size = Pt(9)
        if ci == 1:
            run.bold = True
            v = val.strip().upper()
            if 'HIGH' in v:
                run.font.color.rgb = HIGH_RED
            elif 'MEDIUM' in v:
                run.font.color.rgb = MED_ORANGE
            elif 'LOW' in v:
                run.font.color.rgb = LOW_GREEN
            shade_cell(cell, 'FFFFFF')
            add_cell_borders(cell, color='CCCCCC', sz=4)
        else:
            run.font.color.rgb = BLACK
            shade_cell(cell, 'FFFFFF')
            add_cell_borders(cell, color='CCCCCC', sz=4)

doc.add_paragraph('')




section_heading(doc, '5. Relevance Synthesis')

strengths_header = doc.add_paragraph()
strengths_header.paragraph_format.space_before = Pt(2)
strengths_header.paragraph_format.space_after  = Pt(2)
r = strengths_header.add_run('Strengths of this opportunity:')
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY
strengths = [
    'Defence sector reference — high strategic value for North African market',
    'Full datacenter scope in a single contract — visibility across full ICT portfolio',
    'Specs map cleanly to standard Tier-1 OEM portfolios (no exotic requirements)',
    '2-year support contract generates predictable recurring revenue',
]
for s in strengths:
    bullet_para(doc, s)

doc.add_paragraph('')
challenges_header = doc.add_paragraph()
challenges_header.paragraph_format.space_before = Pt(2)
challenges_header.paragraph_format.space_after  = Pt(2)
r = challenges_header.add_run('Challenges:')
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY
challenges = [
    'VMware migration is the critical-path technical risk',
    'Same-brand constraint requires early OEM commitment',
    'Requires up-to-date hardware (latest generation servers, all-NVMe storage)',
    'Two certifying trainings require an accredited center and certified trainer availability',
]
for c in challenges:
    bullet_para(doc, c)

doc.add_paragraph('')
relevance_header = doc.add_paragraph()
relevance_header.paragraph_format.space_before = Pt(2)
relevance_header.paragraph_format.space_after  = Pt(2)
r = relevance_header.add_run('Relevance by vendor profile:')
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BODY_GREY

rel_table = doc.add_table(rows=0, cols=2)
rel_table.style = 'Table Grid'
add_table_row(rel_table, ['Profile', 'Relevance'], header=True)

profiles = [
    ('Tier-1 OEM with strong North African channel (e.g., Huawei NA)',
     'Highly relevant — natural fit'),
    ('Tier-2 OEM with limited local presence',
     'Moderately relevant — partnership required'),
    ('Local integrator without OEM-level partnership',
     'Low relevance'),
    ('New entrant without local references',
     'Low relevance'),
]
for profile, relevance in profiles:
    add_table_row(rel_table, [profile, relevance])

doc.add_paragraph('')


callout = doc.add_paragraph()
callout.paragraph_format.space_before = Pt(4)
callout.paragraph_format.space_after  = Pt(6)
add_paragraph_border(callout, color='1F3864', sz=6)

run1 = callout.add_run(
    'This tender is technically relevant for any Tier-1 ICT vendor with active North African presence. '
    'The technical scope is standard, the requirements map to common product families, and the reference '
    'value is high. Two technical items require internal validation by management before pursuing: '
    'OEM partnership covering servers + storage + backup, and availability of certified migration profiles.')
run1.font.size = Pt(10)
run1.font.color.rgb = BODY_GREY


import sys
out_path = sys.argv[1] if len(sys.argv) > 1 else '/Users/albus/.openclaw/workspace/offres/analyses/mutuelle-armee-datacenter-2026-05/CCTP_Analysis_Mutuelle_Armee.docx'
doc.save(out_path)
print(f'Saved: {out_path}')

"""
step1_parse.py

Scans input/ for PDF, DOCX, and DOC files, converts each to clean Markdown.
Legacy .doc files are auto-converted to .docx via Microsoft Word COM or LibreOffice.
using Docling (free, local, no API key required), then enriches any
<!-- image --> placeholders by sending the embedded images to Mistral OCR.

Pass 1 — Docling  : text + tables → Markdown (offline, free)
Pass 2 — Mistral  : embedded diagram images → Markdown (API, needs MISTRAL_API_KEY)

If MISTRAL_API_KEY is not set, pass 2 is silently skipped and <!-- image -->
placeholders are kept as-is.

First run: Docling downloads AI models (~1 GB). All subsequent runs are offline.
Run:  python step1_parse.py
"""

import base64
import os
import re
import subprocess
import sys
import time
from pathlib import Path
from typing import Optional

try:
    from dotenv import load_dotenv
    load_dotenv(Path(os.environ.get("RFP_ENV_FILE", Path.home() / ".tuneps_data" / "rfp-pipeline" / ".env")))
except ImportError:
    pass  # python-dotenv optional for pass 1 only

try:
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.datamodel.base_models import InputFormat
except ImportError:
    print("ERROR: docling not installed.")
    print("Run:  pip install docling")
    sys.exit(1)

INPUT_DIR  = Path("input")
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

MIN_IMAGE_PX = 100   # skip images smaller than this in either dimension (icons/logos)


# ─── Pass 1: Docling ──────────────────────────────────────────────────────────

def build_converter() -> DocumentConverter:
    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = False          # text-layer PDF — skip image OCR
    pipeline_options.do_table_structure = True

    return DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
        }
    )


def _safe(s: str) -> str:
    return s.encode(sys.stdout.encoding, errors="replace").decode(sys.stdout.encoding)


def parse_document(file_path: Path, converter: DocumentConverter) -> str:
    """Convert a single PDF or DOCX to Markdown via Docling."""
    print(_safe(f"\nParsing : {file_path.name}"))
    print(f"Size    : {file_path.stat().st_size / 1024:.1f} KB")

    result = converter.convert(str(file_path))
    markdown = result.document.export_to_markdown()

    try:
        print(f"Pages   : {len(result.document.pages)}")
    except Exception:
        pass
    print(f"Chars   : {len(markdown):,}")

    return markdown


def looks_like_heading(line: str) -> bool:
    s = line.strip()
    if not s or len(s) > 120:
        return False
    if 15 < len(s) < 90 and s.isupper() and any(c.isalpha() for c in s):
        return True
    # Real section numbering: 1) Intro, 1.2. Title, 1. Serveur...
    # Avoid treating table/spec values like "2 x alimentations" as headings.
    if re.match(r"^\d+(?:\.\d+)*[\).]\s+\D.{2,}$", s):
        return True
    return False


def fast_parse_pdf(file_path: Path) -> Optional[str]:
    """Fast path for text-layer PDFs using PyMuPDF. Falls back when text is sparse."""
    try:
        import fitz
    except ImportError:
        return None

    t0 = time.time()
    doc = fitz.open(str(file_path))
    parts = []
    for page in doc:
        lines = []
        for raw in page.get_text("text").splitlines():
            s = raw.strip()
            if not s:
                continue
            lines.append(f"## {s}" if looks_like_heading(s) else s)
        if lines:
            parts.append("\n".join(lines))
    pages = len(doc)
    doc.close()

    markdown = "\n\n".join(parts).strip()
    if len(markdown) < 1000:
        return None
    print(_safe(f"\nFast parsing : {file_path.name}"))
    print(f"Size         : {file_path.stat().st_size / 1024:.1f} KB")
    print(f"Pages        : {pages}")
    print(f"Chars        : {len(markdown):,}")
    print(f"Fast parse   : {time.time() - t0:.2f}s")
    return markdown


def fast_parse_docx(file_path: Path) -> Optional[str]:
    """Fast DOCX extraction using python-docx, including table cell text."""
    try:
        import docx
    except ImportError:
        return None

    t0 = time.time()
    document = docx.Document(str(file_path))
    parts = []
    for para in document.paragraphs:
        s = para.text.strip()
        if not s:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or looks_like_heading(s):
            parts.append(f"## {s}")
        else:
            parts.append(s)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    markdown = "\n\n".join(parts).strip()
    if len(markdown) < 1000:
        return None
    print(_safe(f"\nFast parsing : {file_path.name}"))
    print(f"Size         : {file_path.stat().st_size / 1024:.1f} KB")
    print(f"Chars        : {len(markdown):,}")
    print(f"Fast parse   : {time.time() - t0:.2f}s")
    return markdown


def parse_document_auto(file_path: Path, converter: Optional[DocumentConverter]) -> str:
    mode = os.getenv("RFP_PARSE_MODE", "fast").strip().lower()
    if mode in ("fast", "auto"):
        if file_path.suffix.lower() == ".pdf":
            text = fast_parse_pdf(file_path)
        elif file_path.suffix.lower() == ".docx":
            text = fast_parse_docx(file_path)
        else:
            text = None
        if text:
            return text
        if mode == "fast":
            print("WARNING: fast parser produced insufficient text; falling back to Docling")

    if converter is None:
        print("Initialising Docling (fallback parse)...")
        converter = build_converter()
    return parse_document(file_path, converter)


# ─── Pass 2: Mistral OCR for embedded images ─────────────────────────────────

def extract_pdf_images(pdf_path: Path) -> list[dict]:
    """
    Return all images from the PDF that are large enough to be real diagrams.
    Each dict: {bytes, ext, width, height}
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("WARNING: pymupdf not installed — skipping image extraction")
        return []

    doc = fitz.open(str(pdf_path))
    seen_xrefs: set[int] = set()
    images: list[dict] = []

    for page in doc:
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)

            raw = doc.extract_image(xref)
            w, h = raw["width"], raw["height"]
            if w < MIN_IMAGE_PX or h < MIN_IMAGE_PX:
                continue  # skip icons / decorative images

            images.append({
                "bytes": raw["image"],
                "ext":   raw["ext"],
                "width": w,
                "height": h,
            })

    doc.close()
    return images


VISION_PROMPT = (
    "This is a network/IT architecture diagram from a public procurement document. "
    "Describe it completely in structured Markdown:\n"
    "1. List every labeled component (servers, switches, firewalls, storage arrays, etc.) "
    "with its exact label text and quantity if visible.\n"
    "2. Describe every connection between components (which device connects to which, "
    "and the link type/speed if shown).\n"
    "3. Include any legend entries.\n"
    "Be precise and exhaustive — this description replaces the image in a text pipeline."
)


def describe_diagram(image: dict, client) -> str:
    """Send one diagram image to Mistral vision chat; return structured Markdown description."""
    ext = image["ext"].lower()
    mime = f"image/{ext}" if ext in ("png", "jpeg", "jpg", "webp") else "image/png"
    b64 = base64.b64encode(image["bytes"]).decode()
    data_uri = f"data:{mime};base64,{b64}"

    response = client.chat.complete(
        model="pixtral-large-latest",
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": data_uri}},
                    {"type": "text",      "text": VISION_PROMPT},
                ],
            }
        ],
    )
    return response.choices[0].message.content.strip()


def extract_docx_images(docx_path: Path) -> list[dict]:
    """Extract embedded images from a DOCX (zip) file."""
    import io
    import zipfile
    SUPPORTED = {"png", "jpeg", "jpg", "webp", "gif"}
    images: list[dict] = []
    try:
        from PIL import Image as PilImage
        pil_available = True
    except ImportError:
        pil_available = False

    with zipfile.ZipFile(docx_path) as z:
        media = [n for n in z.namelist()
                 if n.startswith("word/media/") and not n.endswith("/")]
        for name in media:
            ext = Path(name).suffix.lstrip(".").lower()
            if ext not in SUPPORTED:
                continue  # skip EMF/WMF vector formats Mistral can't read
            data = z.read(name)
            if pil_available:
                try:
                    img = PilImage.open(io.BytesIO(data))
                    w, h = img.size
                    if w < MIN_IMAGE_PX or h < MIN_IMAGE_PX:
                        continue
                except Exception:
                    continue
            else:
                if len(data) < 5000:
                    continue  # skip tiny blobs without PIL to measure them
                w, h = 999, 999
            images.append({"bytes": data,
                            "ext": "jpeg" if ext == "jpg" else ext,
                            "width": w, "height": h})
    return images


def enrich_with_mistral_ocr(markdown: str, file_path: Path) -> str:
    """
    Replace each <!-- image --> placeholder with Mistral OCR output.
    Falls back to keeping the placeholder on any per-image error.
    Skips entirely if MISTRAL_API_KEY is not configured.
    Supports both PDF and DOCX source files.
    """
    placeholder_count = markdown.count("<!-- image -->")
    if placeholder_count == 0:
        return markdown

    api_key = os.getenv("MISTRAL_API_KEY", "").strip()
    if not api_key:
        print("WARNING: MISTRAL_API_KEY not set — keeping <!-- image --> placeholders")
        return markdown

    ext = file_path.suffix.lower()
    if ext == ".pdf":
        images = extract_pdf_images(file_path)
    elif ext == ".docx":
        images = extract_docx_images(file_path)
    else:
        return markdown

    if not images:
        print(f"WARNING: No extractable images found in {file_path.suffix.upper()}")
        return markdown

    total_imgs   = len(images)
    to_process   = min(placeholder_count, total_imgs)
    print(f"\nImage enrichment: {placeholder_count} placeholder(s), "
          f"{total_imgs} image(s) in PDF -> processing {to_process}")

    try:
        from mistralai.client import Mistral
    except ImportError:
        print("WARNING: mistralai not installed -- skipping image OCR")
        return markdown

    client = Mistral(api_key=api_key)
    ok = 0
    result = markdown

    for i in range(to_process):
        img = images[i]
        print(f"  Processing image {i + 1}/{to_process} "
              f"({img['width']}x{img['height']} px)...", end=" ", flush=True)
        try:
            ocr_text = describe_diagram(img, client)
            if ocr_text.strip():
                replacement = f"[DIAGRAM START]\n{ocr_text.strip()}\n[DIAGRAM END]"
            else:
                replacement = "<!-- image: no text detected -->"
            result = result.replace("<!-- image -->", replacement, 1)
            print("done")
            ok += 1
        except Exception as exc:
            print(f"failed ({exc})")

        if i < to_process - 1:
            time.sleep(1)  # avoid rate-limiting

    skipped = placeholder_count - to_process
    print(f"  Summary: {ok} OCR'd, {to_process - ok} failed, {skipped} skipped")
    return result


# ─── Output ───────────────────────────────────────────────────────────────────

def save_result(text: str, source_file: Path) -> Path:
    out_path = OUTPUT_DIR / (source_file.stem + "_parsed.md")
    out_path.write_text(text, encoding="utf-8")
    print(_safe(f"Saved   : {out_path}"))
    return out_path


# ─── .doc → .docx conversion ─────────────────────────────────────────────────

def convert_doc_to_docx(doc_path: Path) -> Optional[Path]:
    """
    Convert a legacy .doc file to .docx in-place (same folder).
    Tries Microsoft Word COM first (Windows + Word installed),
    then LibreOffice headless as fallback.
    Returns the .docx Path on success, None on failure.
    """
    docx_path = doc_path.with_suffix(".docx")

    # Method 1: Microsoft Word COM automation
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(str(doc_path.resolve()))
        doc.SaveAs2(str(docx_path.resolve()), FileFormat=16)  # 16 = wdFormatDocumentDefault (.docx)
        doc.Close(False)
        word.Quit()
        print(f"  Converted via Microsoft Word → {docx_path.name}")
        return docx_path
    except Exception:
        pass

    # Method 2: LibreOffice headless
    for soffice in ("soffice", r"C:\Program Files\LibreOffice\program\soffice.exe"):
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "docx",
                 "--outdir", str(doc_path.parent), str(doc_path)],
                capture_output=True, timeout=60,
            )
            if result.returncode == 0 and docx_path.exists():
                print(f"  Converted via LibreOffice → {docx_path.name}")
                return docx_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    print(f"  WARNING: Cannot convert {doc_path.name} — install Microsoft Word or LibreOffice")
    return None


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    doc_files = sorted(INPUT_DIR.glob("*.doc"))
    if doc_files:
        print(f"Found {len(doc_files)} legacy .doc file(s) — converting to .docx ...")
        for doc_path in doc_files:
            print(_safe(f"  Converting: {doc_path.name}"))
            convert_doc_to_docx(doc_path)
        print()

    files = sorted(
        list(INPUT_DIR.glob("*.pdf")) + list(INPUT_DIR.glob("*.docx"))
    )

    if not files:
        print(f"No PDF or DOCX files found in {INPUT_DIR}/")
        print("Add files to input/ and run again.")
        sys.exit(0)

    print(f"Found {len(files)} file(s) to parse")
    print(f"Parse mode: {os.getenv('RFP_PARSE_MODE', 'fast')} (fast text extraction with Docling fallback)\n")

    converter = None

    ok, failed = 0, 0
    for file_path in files:
        try:
            # Pass 1: fast text extraction when possible, Docling fallback when needed
            text = parse_document_auto(file_path, converter)
            if not text.strip():
                print(_safe(f"WARNING: {file_path.name} produced empty output"))
                failed += 1
                continue

            # Pass 2: Mistral OCR for embedded images (PDF and DOCX)
            if file_path.suffix.lower() in (".pdf", ".docx"):
                text = enrich_with_mistral_ocr(text, file_path)

            save_result(text, file_path)

            print("\n--- Preview (first 500 chars) ---")
            print(text[:500].encode(sys.stdout.encoding, errors="replace").decode(sys.stdout.encoding))
            print("---")
            ok += 1

        except Exception as e:
            print(_safe(f"ERROR parsing {file_path.name}: {e}"))
            import traceback; traceback.print_exc()
            failed += 1

    print(f"\nDone: {ok} succeeded, {failed} failed")


if __name__ == "__main__":
    main()
